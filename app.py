from flask import Flask, request, jsonify, render_template, redirect, url_for, send_file
import os
import fitz  # For PDF handling
import requests
import json
import openai

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

OLLAMA_ENDPOINT = "http://localhost:11434/api/generate"
HEADERS = {"Content-Type": "application/json"}
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("No Gemini API key found. Please set the GEMINI_API_KEY environment variable.")

# Home Route
@app.route('/')
def home():
    return render_template('index.html')

# Upload Route
@app.route('/upload', methods=['GET', 'POST'])
def upload_resume():
    if request.method == 'POST':
        if 'resume' not in request.files:
            return 'No file uploaded', 400
        file = request.files['resume']
        if file.filename == '':
            return 'No selected file', 400
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)
        content = extract_content(filepath)
        return render_template('edit_resume.html', content=content, original_format=file.filename.split('.')[-1])
    return render_template('upload.html')

@app.route("/resume", methods=["POST"])
def create_resume():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    filename = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filename)
    generated_resume = generate_resume(filename)
    generated_resume_path = os.path.join(app.config['UPLOAD_FOLDER'], 'generated_resume.txt')
    with open(generated_resume_path, 'w') as f:
        f.write(generated_resume)
    return jsonify({"message": "File uploaded successfully", "filename": file.filename, "generated_resume": generated_resume_path}), 200

# Enhance Resume Route
@app.route('/enhance', methods=['POST'])
def enhance_resume():
    content = request.form.get('content')
    ai_service = request.form.get('ai_service', 'ollama')
    objective = request.form.get('objective', None)
    original_format = request.form.get('original_format', 'pdf')
    
    if not content:
        return "Error: No content provided", 400
    
    enhanced_content = enhance_with_ai(content, ai_service, objective)
    return render_template('edit_resume.html', content=enhanced_content, original_format=original_format)

# Download Resume Route
@app.route('/download', methods=['POST'])
def download_resume():
    content = request.form['content']
    original_format = request.form.get('original_format', 'pdf')
    output_filepath = os.path.join(app.config['UPLOAD_FOLDER'], f'enhanced_resume.{original_format}')
    
    if original_format == 'pdf':
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        c = canvas.Canvas(output_filepath, pagesize=letter)
        width, height = letter
        textobject = c.beginText(40, height - 50)
        for line in content.split('\n'):
            textobject.textLine(line)
        c.drawText(textobject)
        c.save()
    elif original_format == 'docx':
        from docx import Document
        doc = Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_filepath)
    else:
        with open(output_filepath, 'w') as f:
            f.write(content)
    
    return send_file(output_filepath, as_attachment=True)

def generate_resume(file_path):
    # Extract text from PDF
    doc = fitz.open(file_path)
    content = ""
    for page in doc:
        content += page.get_text()

    # Call LLM to parse content into neat sections
    payload = {
        "model": "llama3.2",
        "prompt": f"Parse the following resume content into neat sections:\n\n{content}"
    }
    response = requests.post(OLLAMA_ENDPOINT, headers=HEADERS, json=payload)
    response_data = response.json()
    generated_resume = response_data['text']

    return generated_resume

def extract_content(filepath):
    if filepath.endswith('.pdf'):
        # Extract text from PDF
        doc = fitz.open(filepath)
        content = ''
        for page in doc:
            content += page.get_text()
        sections = content.split('\n\n')  # Simple split, adjust as needed
        return sections
    elif filepath.endswith('.docx'):
        from docx import Document
        doc = Document(filepath)
        sections = []
        current_section = ''
        for para in doc.paragraphs:
            text = para.text.strip()
            if text == '':
                if current_section:
                    sections.append(current_section)
                    current_section = ''
            else:
                current_section += text + '\n'
        if current_section:
            sections.append(current_section)
        return sections
    else:
        return []

def enhance_with_ai(content, ai_service, objective=None):
    if ai_service == "ollama":
        return enhance_with_ollama(content, objective)
    elif ai_service == "gemini":
        return enhance_with_gemini(content, objective)
    elif ai_service == "openai":
        return enhance_with_openai(content, objective)
    else:
        return "Error: Unsupported AI service selected."

def enhance_with_ollama(content, objective=None):
    prompt = "You are a professional resume writer.\n"
    if objective:
        prompt += f"Objective: {objective}\n"
    prompt += f"Enhance the following resume content to make it more effective and impactful:\n\n{content}"
    
    payload = {
        "model": "llama3.2",  # Replace with the actual model name if different
        "prompt": prompt
    }
    try:
        response = requests.post(OLLAMA_ENDPOINT, headers=HEADERS, json=payload)
        response.raise_for_status()  # Raise an error for bad status codes
        response_data = response.json()
        return response_data.get('text', content)
    except requests.exceptions.RequestException as e:
        print(f"Ollama API Error: {e}")
        return "Error: Failed to enhance resume using Ollama."
    except ValueError:
        print("Ollama API returned non-JSON response.")
        return "Error: Invalid response from Ollama."

def enhance_with_gemini(content, objective=None):
    GEMINI_ENDPOINT = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent"
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")  # Ensure this is set

    if not GEMINI_API_KEY:
        print("Gemini API key is not set.")
        return "Error: Gemini API key is missing."

    headers = {
        "Content-Type": "application/json",
    }
    prompt = "You are a professional resume writer.\n"
    if objective:
        prompt += f"Objective: {objective}\n"
    prompt += f"Enhance the following resume content to make it more effective and impactful:\n\n{content}"
    payload = {
        "contents": [
            {
                "parts": [
                    {
                        "text": prompt
                    }
                ]
            }
        ]
    }

    params = {
        "key": GEMINI_API_KEY
    }

    try:
        response = requests.post(GEMINI_ENDPOINT, headers=headers, params=params, json=payload)
        response.raise_for_status()  # Raise an error for bad status codes
        response_data = response.json()
        print(json.dumps(response_data['candidates'][0]['content']))
        
        # Parse the response based on Gemini's API structure
        enhanced_text = ""
        if 'content' in response_data.get('candidates', [{}])[0]:
            for part in response_data['candidates'][0]['content']['parts']:
                enhanced_text += part.get('text', '')
        
        return enhanced_text if enhanced_text else "Error: No enhanced content received from Gemini."

    except requests.exceptions.RequestException as e:
        print(f"Gemini API Error: {e}")
        return "Error: Failed to enhance resume using Gemini."
    except ValueError as e:
        print(f"Invalid response from Gemini: {e}")
        return "Error: Invalid response from Gemini."

def enhance_with_openai(content, objective=None):
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    if not OPENAI_API_KEY:
        print("OpenAI API key is not set.")
        return "Error: OpenAI API key is missing."
    openai.api_key = OPENAI_API_KEY

    prompt = "You are a professional resume writer.\n"
    if objective:
        prompt += f"Objective: {objective}\n"
    prompt += f"Enhance the following resume content to make it more effective and impactful:\n\n{content}"

    try:
        response = openai.Completion.create(
            engine="text-davinci-003",
            prompt=prompt,
            max_tokens=1024,
            temperature=0.7,
            n=1,
            stop=None,
        )
        enhanced_text = response.choices[0].text.strip()
        return enhanced_text if enhanced_text else "Error: No enhanced content received from OpenAI."

    except openai.error.OpenAIError as e:
        print(f"OpenAI API Error: {e}")
        return "Error: Failed to enhance resume using OpenAI."

if __name__ == "__main__":
    app.run(debug=True)
